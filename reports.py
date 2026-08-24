"""
ABLLS-R 评估报告生成模块
- 生成维度得分柱状图 (matplotlib, 中文)
- 生成 Word 报告 (python-docx)
- 生成 PDF 报告 (reportlab)
"""

import io
import os
import base64
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager

# 中文字体配置：优先项目内子集字体，再尝试 Windows / Linux 系统字体
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CN_FONT_CANDIDATES = [
    os.path.join(_BASE_DIR, "assets", "fonts", "SubsetSourceHanSansSC-Regular.otf"),
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
]


def _find_cn_font():
    for path in CN_FONT_CANDIDATES:
        if os.path.exists(path):
            try:
                font_manager.fontManager.addfont(path)
                name = font_manager.FontProperties(fname=path).get_name()
                return path, name
            except Exception:
                continue
    return None, None


FONT_PATH, FONT_NAME = _find_cn_font()
if FONT_NAME:
    plt.rcParams["font.family"] = FONT_NAME
plt.rcParams["axes.unicode_minus"] = False

C_BLUE = "#2f6fed"
C_GREEN = "#1faa59"
C_GREY = "#cbd5e1"

# 与 app.py 中 ACHIEVE_THRESHOLD 保持一致：得分率 ≥ 该值视为基本掌握
ACHIEVE_THRESHOLD_RPT = 0.8
# 基础（前提）领域：作为后续技能的基础，排序时自动加权
FOUNDATIONAL_DOMAINS = {
    "配合与强化物的效能", "视觉表现", "接受性语言", "模仿", "命名、描述",
}

def age_label_rpt(months):
    if months is None:
        return "—"
    y = months // 12
    m = months % 12
    if y == 0:
        return f"{m}个月"
    if m == 0:
        return f"{y}岁"
    return f"{y}岁{m}个月"


def _frac(score, mx):
    return score / mx if mx > 0 else 0.0


def _short_suggestion(state, domain):
    """状态·教学建议列：状态 + 教学方法缩写"""
    if state.startswith("未掌握"):
        return "DTT"
    # 临界
    return "NET"


def make_bar_chart(domains, path=None):
    """生成各维度得分率横向柱状图，返回图片字节（path 给定则存文件）"""
    ds = [d for d in domains if d["max_score"] > 0]
    ds = sorted(ds, key=lambda x: x["rate"])
    names = [d["domain"] for d in ds]
    rates = [d["rate"] for d in ds]
    fig, ax = plt.subplots(figsize=(7.2, max(4, len(names) * 0.42)))
    colors = [C_GREEN if r >= 70 else (C_BLUE if r >= 40 else "#e2554f") for r in rates]
    ax.barh(names, rates, color=colors)
    for i, r in enumerate(rates):
        ax.text(min(r + 1, 96), i, f"{r}%", va="center", fontsize=8)
    ax.set_xlim(0, 100)
    ax.set_xlabel("得分率 (%)")
    ax.set_title("各维度得分率", fontsize=12)
    ax.grid(axis="x", linestyle=":", alpha=0.5)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130)
    plt.close(fig)
    data = buf.getvalue()
    if path:
        with open(path, "wb") as f:
            f.write(data)
    return data


def make_band_chart(domains, path=None):
    """生成领域发展层级分布图（良好/中等/薄弱），返回图片字节"""
    active = [d for d in domains if d["max_score"] > 0 and d["rated"] > 0]
    bands = ["良好(≥70%)", "中等(40-69%)", "薄弱(<40%)"]
    counts = [0, 0, 0]
    for d in active:
        if d["rate"] >= 70:
            counts[0] += 1
        elif d["rate"] >= 40:
            counts[1] += 1
        else:
            counts[2] += 1
    fig, ax = plt.subplots(figsize=(5.6, 3.2))
    bcolors = [C_GREEN, C_BLUE, "#e2554f"]
    bars = ax.bar(bands, counts, color=bcolors)
    for b, v in zip(bars, counts):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.05, str(v),
                ha="center", va="bottom", fontsize=11, fontweight="bold")
    ax.set_ylabel("领域数")
    ax.set_title("领域发展层级分布", fontsize=12)
    ax.set_ylim(0, max(counts) + 1 if any(counts) else 1)
    ax.grid(axis="y", linestyle=":", alpha=0.5)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130)
    plt.close(fig)
    data = buf.getvalue()
    if path:
        with open(path, "wb") as f:
            f.write(data)
    return data


def band_label(rate):
    if rate >= 70:
        return "良好"
    if rate >= 40:
        return "中等"
    return "薄弱"


def is_main_item_rpt(age, adapt):
    """项月龄为空、或适配月龄为空、或项月龄 <= 适配月龄 => 主评估；
    超出适配月龄的为拓展评估，不计入干预清单。"""
    return (age is None) or (adapt is None) or (age <= adapt)


def build_chart_analysis(domains, total):
    """生成图表的结构化解读，返回 [('h'|'p', 文本), ...]"""
    active = [d for d in domains if d["max_score"] > 0 and d["rated"] > 0]
    out = []
    if not active:
        return [("p", "本次评估暂无可解读的数据，建议完成至少一个领域的评分后再生成解读。")]
    by_rate = sorted(active, key=lambda x: x["rate"])
    total_rate = total["rate"]
    out.append(("p",
        f"本次评估总得分率为 {total_rate}%（{total['score']}/{total['max_total']} 分），"
        f"下方图表共展示 {len(active)} 个已评领域，按得分率由低到高排列，并以颜色区分发展层级："
        f"绿色（≥70%）表示发展良好，蓝色（40%–69%）表示部分掌握、需巩固，"
        f"红色（<40%）表示明显薄弱、需优先干预。"))
    # 优势领域
    top = by_rate[-3:][::-1]
    strengths = "；".join(f"{d['domain']}（{d['rate']}%）" for d in top)
    out.append(("h", "1. 优势领域"))
    out.append(("p",
        f"得分率最高的领域为：{strengths}。这些领域可视为儿童的相对优势，"
        f"后续教学可在其基础上衔接新技能，借助已掌握的沟通能力提升参与度与成功率。"))
    # 需优先加强
    weak = [d for d in by_rate if d["rate"] < 70]
    out.append(("h", "2. 需优先加强的领域"))
    if weak:
        bottom = weak[:3]
        weak_txt = "；".join(f"{d['domain']}（{d['rate']}%）" for d in bottom)
        out.append(("p",
            f"得分率偏低、建议优先安排密集干预的领域为：{weak_txt}。"
            f"其中红色条（<40%）代表目前几乎尚未掌握，宜将目标技能拆解为更小步骤、"
            f"增加示范与肢体辅助，并配合高频强化逐步建立。"))
    else:
        out.append(("p",
            "当前已评领域得分率均在 70% 以上，暂未见明显薄弱项，建议保持现有教学节奏并定期复评。"))
    # 解读说明（覆盖度）
    out.append(("h", "3. 解读说明"))
    unrated = total["total_items"] - total["rated"]
    if unrated > 0:
        out.append(("p",
            f"本次共评估 {total['rated']}/{total['total_items']} 项，未评估的 {unrated} 项按 0 分计入得分率，"
            f"因此部分领域得分率偏低可能源于评估覆盖不全，建议补全评估后再做最终判断。"))
    else:
        out.append(("p", "本次已覆盖全部评估项目，得分率可反映真实能力基线。"))
    return out


def build_report_data(aid, db_get):
    """从数据库收集生成报告所需的全部数据"""
    a = db_get(aid)
    client_id = a["client_id"]
    return a, client_id


def generate_docx(assessment, client, total, domains, charts, analysis, ability, suggestions):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 标题
    h = doc.add_heading("语言与学习技能评估报告", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    doc.add_paragraph()
    info = doc.add_table(rows=0, cols=2)
    info.style = "Light List Accent 1"
    rows = [
        ("被评估者", client.get("name", "")),
        ("性别", client.get("gender", "") or "—"),
        ("出生日期", client.get("birth_date", "") or "—"),
        ("评估标题", assessment.get("title", "")),
        ("评估员", assessment.get("assessor", "") or "—"),
        ("评估日期", assessment.get("date", "")),
        ("报告生成时间", assessment.get("created_at", "")),
    ]
    for k, v in rows:
        c = info.add_row().cells
        c[0].text = k
        c[1].text = str(v)
        c[0].paragraphs[0].runs[0].bold = True

    # 总分概览
    doc.add_heading("一、总体得分概览", level=1)
    p = doc.add_paragraph()
    p.add_run("总分：").bold = True
    p.add_run(f"{total['score']} / {total['max_total']} 分（得分率 {total['rate']}%）")
    p2 = doc.add_paragraph()
    p2.add_run("已评估项目：").bold = True
    p2.add_run(f"{total['rated']} / {total['total_items']} 项")

    # 维度得分表
    doc.add_heading("二、各维度得分明细", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, txt in enumerate(["领域", "得分", "满分", "得分率", "已评/总项"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
    for d in domains:
        c = t.add_row().cells
        c[0].text = d["domain"]
        c[1].text = str(d["score"])
        c[2].text = str(d["max_score"])
        c[3].text = f"{d['rate']}%"
        c[4].text = f"{d['rated']}/{d['item_count']}"

    # 图表
    doc.add_heading("三、各维度得分率可视化", level=1)
    if charts:
        for i, cb in enumerate(charts):
            if cb:
                doc.add_picture(io.BytesIO(cb), width=Inches(6.0))
                if i == 0:
                    doc.add_paragraph("图1：各领域得分率（由低到高排列，颜色区分发展层级）")
                else:
                    doc.add_paragraph("图2：领域发展层级分布（按得分率归入良好/中等/薄弱）")

    # 图表解读
    doc.add_heading("四、图表解读", level=1)
    for kind, txt in (analysis or []):
        if kind == "h":
            doc.add_heading(txt, level=2)
        else:
            doc.add_paragraph(txt)

    # 干预建议（结构化：优先度总览 + 分领域具体项目）
    doc.add_heading("五、干预建议（基于本次评估）", level=1)
    if not suggestions:
        doc.add_paragraph("本次评估各已评领域均达到实际年龄水平，暂无明显需优先干预的薄弱领域，建议保持现有教学节奏并定期复评。")
    else:
        doc.add_paragraph(
            "以下按「干预优先度」排序生成。优先度由三项合并计算：① 能力年龄滞后实际年龄的月数"
            "（落后越多越优先）；② 基础领域（配合、视觉、接受性语言、模仿、命名）自动加权；"
            "③ 该领域「临界可突破项」数量（已部分掌握、再教即可突破，投入产出比高者加权）。"
            "P1 为最高优先，依次递减。"
        )
        doc.add_paragraph(
            "说明：本干预清单仅列入「主评估（适龄）窗口」内的未达标项；拓展评估（超出适配月龄）的"
            "超龄技能不计入，以免目标偏离孩子当前发育阶段。"
        )
        # 优先度总览表
        t = doc.add_table(rows=1, cols=6)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, txt in enumerate(["优先度", "领域", "能力年龄", "实际年龄", "滞后(月)", "未达标项(临界)"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
        for s in suggestions:
            c = t.add_row().cells
            c[0].text = s["tier"]
            c[1].text = s["domain"]
            c[2].text = s["ability_label"]
            c[3].text = s["chron_label"]
            c[4].text = "—" if s["lag_proxy"] else str(round(s["lag"]))
            c[5].text = f"{s['unmastered_count']}（{s['critical']}）"
        # 分领域明细
        for s in suggestions:
            doc.add_heading(f"{s['tier']} · {s['domain']}", level=2)
            p = doc.add_paragraph()
            p.add_run("干预理由：").bold = True
            p.add_run(s["reason"])
            if s["items"]:
                t2 = doc.add_table(rows=1, cols=6)
                t2.style = "Light Grid Accent 1"
                hdr2 = t2.rows[0].cells
                for i, txt in enumerate(["编号", "技能点", "任务", "适用月龄", "当前/满分", "状态 · 教学建议"]):
                    hdr2[i].text = txt
                    hdr2[i].paragraphs[0].runs[0].bold = True
                for it in s["items"]:
                    c = t2.add_row().cells
                    c[0].text = it["code"]
                    c[1].text = it["skill_point"]
                    c[2].text = it["task_name"]
                    c[3].text = age_label_rpt(it["age"])
                    c[4].text = f"{it['score']}/{it['max']}"
                    c[5].text = f"{it['state']} ｜ {it['strategy']}"
            else:
                doc.add_paragraph("（该领域未达标项已在总评中体现，暂无逐项明细。）")

    # 能力年龄评估
    doc.add_heading("六、各维度能力年龄评估", level=1)
    _render_ability_docx(doc, ability)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_ability_docx(doc, ability):
    if not ability:
        doc.add_paragraph("暂无能力年龄数据。")
        return
    chron = ability.get("chronological_label", "—")
    overall = ability.get("overall_label", "—")
    p = doc.add_paragraph()
    p.add_run("实际年龄：").bold = True
    p.add_run(f"{chron}（按评估日与出生日期计算）。")
    p2 = doc.add_paragraph()
    p2.add_run("综合能力年龄（取各维度最低值为下限）：").bold = True
    p2.add_run(f"{overall}。")
    note = doc.add_paragraph()
    note.add_run("判定方法：").bold = True
    thr = int(round((ability.get("achieve_threshold") or 0.8) * 100))
    note.add_run(f"得分率 ≥ {thr}% 为基本掌握，<{thr}% 为临界；取连续掌握最远点线性插值定龄；"
                 f"±3 个月内为符合年龄，超出为超前/滞后。综合能力年龄取各维度最低值。")
    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, txt in enumerate(["领域", "能力年龄", "实际年龄", "结论", "基本掌握/定龄项", "判定说明"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
    for d in ability.get("domains", []):
        c = t.add_row().cells
        c[0].text = d["domain"]
        c[1].text = d["ability_label"]
        c[2].text = d["chronological_label"]
        gap = d.get("gap")
        if d["status"] == "滞后" and gap is not None:
            concl = f"滞后 {abs(gap)} 个月"
        elif d["status"] == "超前" and gap is not None:
            concl = f"超前 {gap} 个月"
        else:
            concl = d["status"]
        c[3].text = concl
        c[4].text = f"{d['achieved']}/{d['scored_with_age']}"
        c[5].text = d.get("note", "")


def generate_pdf(assessment, client, total, domains, charts, analysis, ability, suggestions):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, Image, ListFlowable, ListItem)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    # 注册中文字体
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    CN = "Helvetica"
    if FONT_PATH and os.path.exists(FONT_PATH):
        try:
            pdfmetrics.registerFont(TTFont("CN", FONT_PATH))
            CN = "CN"
        except Exception:
            pass

    styles = getSampleStyleSheet()
    st_title = ParagraphStyle("t", parent=styles["Title"], fontName=CN, fontSize=18)
    st_h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName=CN, fontSize=13)
    st_body = ParagraphStyle("b", parent=styles["BodyText"], fontName=CN, fontSize=10)
    st_cell = ParagraphStyle("c", parent=styles["BodyText"], fontName=CN, fontSize=9)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title="评估报告")
    elems = []
    elems.append(Paragraph("语言与学习技能评估报告", st_title))
    elems.append(Spacer(1, 10))

    def P(txt, bold=False):
        return Paragraph(("<b>%s</b>" % txt) if bold else str(txt), st_cell)

    info_rows = [
        [P("被评估者", True), P(client.get("name", "")), P("性别", True), P(client.get("gender", "") or "—")],
        [P("出生日期", True), P(client.get("birth_date", "") or "—"), P("评估标题", True), P(assessment.get("title", ""))],
        [P("评估员", True), P(assessment.get("assessor", "") or "—"), P("评估日期", True), P(assessment.get("date", ""))],
    ]
    it = Table(info_rows, colWidths=[2.5 * cm, 5 * cm, 2.5 * cm, 5 * cm])
    it.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef3ff")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#eef3ff")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elems.append(it)
    elems.append(Spacer(1, 12))

    elems.append(Paragraph("一、总体得分概览", st_h1))
    elems.append(Paragraph(
        f"总分 <b>{total['score']} / {total['max_total']}</b> 分（得分率 <b>{total['rate']}%</b>）；"
        f"已评估项目 {total['rated']} / {total['total_items']} 项。", st_body))
    elems.append(Spacer(1, 8))

    elems.append(Paragraph("二、各维度得分明细", st_h1))
    data = [[P("领域", True), P("得分", True), P("满分", True), P("得分率", True), P("已评/总项", True)]]
    for d in domains:
        data.append([P(d["domain"]), P(str(d["score"])), P(str(d["max_score"])),
                     P(f"{d['rate']}%"), P(f"{d['rated']}/{d['item_count']}")])
    dt = Table(data, colWidths=[6 * cm, 2 * cm, 2 * cm, 2.2 * cm, 2.8 * cm], repeatRows=1)
    dt.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2f6fed")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    elems.append(dt)
    elems.append(Spacer(1, 12))

    if charts:
        elems.append(Paragraph("三、各维度得分率可视化", st_h1))
        for i, cb in enumerate(charts):
            if cb:
                img = Image(io.BytesIO(cb), width=14.5 * cm, height=14.5 * cm * 0.58)
                elems.append(img)
                cap = ("图1：各领域得分率（由低到高排列，颜色区分发展层级）"
                       if i == 0 else "图2：领域发展层级分布（良好/中等/薄弱）")
                elems.append(Paragraph(cap, st_cell))
                elems.append(Spacer(1, 6))
        elems.append(Spacer(1, 6))

    elems.append(Paragraph("四、图表解读", st_h1))
    for kind, txt in (analysis or []):
        if kind == "h":
            elems.append(Paragraph(("<b>%s</b>" % txt), st_body))
        else:
            elems.append(Paragraph(txt, st_body))
        elems.append(Spacer(1, 4))

    elems.append(Paragraph("五、干预建议（基于本次评估）", st_h1))
    if suggestions:
        elems.append(Paragraph(
            "以下按「干预优先度」排序。优先度由三项合并：① 能力年龄滞后实际年龄月数（落后越多越优先）；"
            "② 基础领域（配合、视觉、接受性语言、模仿、命名）自动加权；③ 该领域「临界可突破项」数量"
            "（再教即可突破，投入产出比高）。P1 最高优先。", st_body))
        elems.append(Paragraph(
            "说明：本干预清单仅列入「主评估（适龄）窗口」内的未达标项；拓展评估（超出适配月龄）的"
            "超龄技能不计入，以免目标偏离孩子当前发育阶段。", st_body))
        elems.append(Spacer(1, 6))
        data = [[P("优先度", True), P("领域", True), P("能力年龄", True),
                 P("实际年龄", True), P("滞后(月)", True), P("未达标项(临界)", True)]]
        for s in suggestions:
            lagtxt = "—" if s["lag_proxy"] else str(round(s["lag"]))
            data.append([P(s["tier"]), P(s["domain"]), P(s["ability_label"]),
                         P(s["chron_label"]), P(str(round(s["lag"])) if not s["lag_proxy"] else "—"),
                         P(f"{s['unmastered_count']}（{s['critical']}）")])
        ot = Table(data, colWidths=[1.6 * cm, 3.2 * cm, 2.2 * cm, 2.2 * cm, 1.8 * cm, 3.0 * cm], repeatRows=1)
        ot.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2f6fed")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elems.append(ot)
        elems.append(Spacer(1, 8))
        for s in suggestions:
            elems.append(Paragraph(f"<b>{s['tier']} · {s['domain']}</b>", st_body))
            elems.append(Paragraph(("<b>干预理由：</b>" + s["reason"]), st_body))
            elems.append(Spacer(1, 3))
            if s["items"]:
                idata = [[P("编号", True), P("技能点", True), P("任务", True),
                          P("适用月龄", True), P("当前/满分", True), P("状态·教学建议", True)]]
                for it in s["items"]:
                    idata.append([P(it["code"]), P(it["skill_point"]), P(it["task_name"]),
                                  P(age_label_rpt(it["age"])), P(f"{it['score']}/{it['max']}"),
                                  P(f"{it['state']} · {it['strategy']}")])
                itt = Table(idata, colWidths=[1.5 * cm, 2.3 * cm, 3.0 * cm, 1.7 * cm, 1.6 * cm, 5.3 * cm], repeatRows=1)
                itt.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2f6fed")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                elems.append(itt)
            else:
                elems.append(Paragraph("（该领域未达标项已在总评中体现，暂无逐项明细。）", st_body))
            elems.append(Spacer(1, 8))
        # 缩写说明
        if suggestions:
            from reportlab.lib.styles import ParagraphStyle
            st_legend = ParagraphStyle("legend", parent=st_body, fontSize=11, textColor=colors.HexColor("#6b7280"))
            elems.append(Paragraph(
                "<b>缩写说明：</b>"
                "DTT = 离散单元教学（Discrete Trial Training），适用于从零建立全新技能；"
                "NET = 自然情境教学（Natural Environment Teaching），适用于已部分掌握的技能进行临界突破。",
                st_legend))
            elems.append(Spacer(1, 4))
    else:
        elems.append(Paragraph("本次评估各已评领域均达到实际年龄水平，暂无明显需优先干预的薄弱领域，建议保持现有教学节奏并定期复评。", st_body))

    # 六、能力年龄评估
    elems.append(Paragraph("六、各维度能力年龄评估", st_h1))
    if ability:
        chron = ability.get("chronological_label", "—")
        overall = ability.get("overall_label", "—")
        elems.append(Paragraph(
            f"实际年龄：<b>{chron}</b>（按评估日与出生日期计算）；"
            f"综合能力年龄（取各维度最低值为下限）：<b>{overall}</b>。", st_body))
        elems.append(Paragraph(
            "判定方法："
            "得分率 ≥ 80% 为基本掌握，<80% 为临界；取连续掌握最远点线性插值定龄；"
            "±3 个月内为符合年龄，超出为超前/滞后。综合能力年龄取各维度最低值。",
            st_body))
        elems.append(Spacer(1, 4))
        data = [[P("领域", True), P("能力年龄", True), P("实际年龄", True),
                 P("结论", True), P("基本掌握/定龄项", True), P("判定说明", True)]]
        for d in ability.get("domains", []):
            gap = d.get("gap")
            if d["status"] == "滞后" and gap is not None:
                concl = f"滞后 {abs(gap)} 个月"
            elif d["status"] == "超前" and gap is not None:
                concl = f"超前 {gap} 个月"
            else:
                concl = d["status"]
            data.append([P(d["domain"]), P(d["ability_label"]), P(d["chronological_label"]),
                         P(concl), P(f"{d['achieved']}/{d['scored_with_age']}"), P(d.get("note", ""))])
        at = Table(data, colWidths=[3.6 * cm, 2.1 * cm, 2.1 * cm, 2.8 * cm, 2.4 * cm, 3.6 * cm], repeatRows=1)
        at.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2f6fed")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elems.append(at)
    else:
        elems.append(Paragraph("暂无能力年龄数据。", st_body))

    doc.build(elems)
    return buf.getvalue()


def build_suggestions(domains, total, ability, scale_items, scored, adapt_months=None):
    """生成结构化干预建议。

    优先度由三项合并计算（复合分越高越优先）：
      ① 能力年龄滞后实际年龄的月数（落后越多越优先）；
      ② 基础领域（配合、视觉、接受性语言、模仿、命名）自动加权；
      ③ 该领域「临界可突破项」数量（已部分掌握、再教即可突破，投入产出比高）。
    对每个需干预领域收集「全部未达标项」并附教学建议，按复合分降序标注 P1/P2/P3。
    """
    chron = ability.get("chronological")
    ab_map = {d["domain"]: d for d in ability.get("domains", [])}
    out = []
    for d in domains:
        dname = d["domain"]
        ab = ab_map.get(dname)
        ability_age = ab.get("ability_age") if ab else None
        # 滞后月数（真实月龄差；无可定龄项时以得分率缺口作代理，仅用于排序）
        if ability_age is not None and chron is not None and ability_age < chron:
            lag = float(chron - ability_age)
            lag_proxy = False
        elif ability_age is None:
            lag = max(0.0, (100 - d["rate"])) / 4.0
            lag_proxy = True
        else:
            lag = 0.0
            lag_proxy = False
        # 收集未达标项
        unmastered = []
        critical = 0
        for code, it in scale_items.items():
            if it.get("domain") != dname:
                continue
            if not is_main_item_rpt(it.get("age"), adapt_months):
                continue  # 拓展评估（超龄）项不列入干预清单
            mx = it.get("max_score") or 0
            sc = scored.get(code)
            if sc is None or mx <= 0 or sc >= mx:
                continue
            f = _frac(sc, mx)
            if sc == 0:
                state = "未掌握"
            elif f < ACHIEVE_THRESHOLD_RPT:
                state = "临界"
                critical += 1
            else:
                continue  # 已掌握不列入
            unmastered.append({
                "code": code,
                "skill_point": it.get("skill_point", ""),
                "task_name": it.get("task_name", ""),
                "age": it.get("age"),
                "score": sc,
                "max": mx,
                "state": state,
                "strategy": _short_suggestion(state, dname),
            })
        if d["rated"] <= 0:
            continue  # 未评估领域不纳入干预优先排序（属"未知"而非"薄弱"）
        if not unmastered:
            continue  # 主评估窗口内无未达标项，不列入干预清单（拓展项已排除）
        fw = 1.3 if dname in FOUNDATIONAL_DOMAINS else 1.0
        composite = lag * fw + critical * 3.0
        # 理由
        if ability_age is None:
            reason = (f"该领域暂无可定龄评分项，无法精确估算能力年龄；但本次仍有 {len(unmastered)} 项未达标，"
                      f"建议补全适龄项目评估后再定级。")
        else:
            reason = f"能力年龄约 {age_label_rpt(ability_age)}，滞后实际年龄约 {round(lag)} 个月"
            if fw > 1:
                reason += "；该领域属基础技能前提，已加权优先"
            if critical > 0:
                reason += f"；含 {critical} 项临界可突破目标（再教即可突破，投入产出比高）"
        out.append({
            "domain": dname,
            "ability_label": age_label_rpt(ability_age),
            "chron_label": age_label_rpt(chron),
            "lag": round(lag, 1),
            "lag_proxy": lag_proxy,
            "rate": d["rate"],
            "critical": critical,
            "unmastered_count": len(unmastered),
            "foundational": fw > 1,
            "composite": round(composite, 1),
            "reason": reason,
            "items": sorted(unmastered, key=lambda x: (x["age"] if x["age"] is not None else 999)),
        })
    if not out:
        return []
    out.sort(key=lambda x: x["composite"], reverse=True)
    for i, s in enumerate(out):
        s["tier"] = "P1" if i < 3 else ("P2" if i < 6 else "P3")
    return out


def generate_html(assessment, client, total, domains, charts, analysis, ability, suggestions):
    """生成可在浏览器内预览的完整 HTML 报告（与 Word/PDF 内容一致），图表以 base64 内嵌。"""
    def esc_h(s):
        return (str(s).replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;").replace('"', "&quot;"))

    def img_b64(b):
        return "data:image/png;base64," + base64.b64encode(b).decode("ascii") if b else ""

    # 基本信息
    info = [
        ("被评估者", client.get("name", "") or "—"),
        ("性别", client.get("gender", "") or "—"),
        ("出生日期", client.get("birth_date", "") or "—"),
        ("评估标题", assessment.get("title", "") or "—"),
        ("评估员", assessment.get("assessor", "") or "—"),
        ("评估日期", assessment.get("date", "") or "—"),
        ("报告生成时间", assessment.get("created_at", "") or "—"),
    ]
    info_rows = "".join(
        f'<tr><th>{esc_h(k)}</th><td>{esc_h(v)}</td></tr>' for k, v in info
    )
    # 二、各维度得分明细
    detail_rows = "".join(
        f'<tr><td>{esc_h(d["domain"])}</td><td>{d["score"]}</td><td>{d["max_score"]}</td>'
        f'<td>{d["rate"]}%</td><td>{d["rated"]}/{d["item_count"]}</td></tr>'
        for d in domains
    )
    # 三、图表
    chart_imgs = ""
    for i, cb in enumerate(charts):
        if cb:
            cap = "图1：各领域得分率（由低到高排列，颜色区分发展层级）" if i == 0 \
                else "图2：领域发展层级分布（按得分率归入良好/中等/薄弱）"
            chart_imgs += f'<div class="chart"><img src="{img_b64(cb)}"/><div class="cap">{cap}</div></div>'
    # 四、图表解读
    analysis_html = ""
    for kind, txt in (analysis or []):
        if kind == "h":
            analysis_html += f"<h3>{esc_h(txt)}</h3>"
        else:
            analysis_html += f"<p>{esc_h(txt)}</p>"
    # 五、干预建议
    if not suggestions:
        sugg_html = ("<p>本次评估各已评领域均达到实际年龄水平，暂无明显需优先干预的薄弱领域，"
                     "建议保持现有教学节奏并定期复评。</p>")
    else:
        overview_rows = "".join(
            f'<tr><td>{s["tier"]}</td><td>{esc_h(s["domain"])}</td><td>{s["ability_label"]}</td>'
            f'<td>{s["chron_label"]}</td><td>{"—" if s["lag_proxy"] else round(s["lag"])}</td>'
            f'<td>{s["unmastered_count"]}（{s["critical"]}）</td></tr>'
            for s in suggestions
        )
        detail_blocks = ""
        for s in suggestions:
            if s["items"]:
                items_rows = "".join(
                    f'<tr><td>{it["code"]}</td><td>{esc_h(it["skill_point"])}</td>'
                    f'<td>{esc_h(it["task_name"])}</td><td>{age_label_rpt(it["age"])}</td>'
                    f'<td>{it["score"]}/{it["max"]}</td>'
                    f'<td>{esc_h(it["state"])} · {esc_h(it["strategy"])}</td></tr>'
                    for it in s["items"]
                )
                items_html = ('<table class="grid"><thead><tr>'
                              '<th>编号</th><th>技能点</th><th>任务</th><th>适用月龄</th>'
                              '<th>当前/满分</th><th>状态 · 教学建议</th>'
                              '</tr></thead><tbody>' + items_rows + '</tbody></table>')
            else:
                items_html = '<p class="muted">（该领域未达标项已在总评中体现，暂无逐项明细。）</p>'
            detail_blocks += (f'<div class="sugg-block"><h4>{esc_h(s["tier"])} · {esc_h(s["domain"])}</h4>'
                              f'<p><b>干预理由：</b>{esc_h(s["reason"])}</p>{items_html}</div>')
        sugg_html = (
            "<p>以下按「干预优先度」排序生成。优先度由三项合并计算：① 能力年龄滞后实际年龄的月数"
            "（落后越多越优先）；② 基础领域（配合、视觉、接受性语言、模仿、命名）自动加权；"
            "③ 该领域「临界可突破项」数量（已部分掌握、再教即可突破，投入产出比高者加权）。"
            "P1 为最高优先，依次递减。</p>"
            "<p class='muted'>说明：本干预清单仅列入「主评估（适龄）窗口」内的未达标项；拓展评估"
            "（超出适配月龄）的超龄技能不计入，以免目标偏离孩子当前发育阶段。</p>"
            '<table class="grid"><thead><tr><th>优先度</th><th>领域</th><th>能力年龄</th>'
            '<th>实际年龄</th><th>滞后(月)</th><th>未达标项(临界)</th></tr></thead><tbody>'
            + overview_rows + '</tbody></table>' + detail_blocks
            + '<p class="muted" style="margin-top:10px;font-size:12px">'
            '<b>缩写说明：</b>'
            'DTT = 离散单元教学（Discrete Trial Training），适用于从零建立全新技能；'
            'NET = 自然情境教学（Natural Environment Teaching），适用于已部分掌握的技能进行临界突破。</p>'
        )
    # 六、能力年龄
    if ability:
        chron = ability.get("chronological_label", "—")
        overall = ability.get("overall_label", "—")
        thr = int(round((ability.get("achieve_threshold") or 0.8) * 100))
        ab_rows = ""
        for d in ability.get("domains", []):
            gap = d.get("gap")
            if d["status"] == "滞后" and gap is not None:
                concl = f"滞后 {abs(gap)} 个月"
            elif d["status"] == "超前" and gap is not None:
                concl = f"超前 {gap} 个月"
            else:
                concl = d["status"]
            ab_rows += (f'<tr><td>{esc_h(d["domain"])}</td><td>{d["ability_label"]}</td>'
                        f'<td>{d["chronological_label"]}</td><td>{esc_h(concl)}</td>'
                        f'<td>{d["achieved"]}/{d["scored_with_age"]}</td>'
                        f'<td>{esc_h(d.get("note", ""))}</td></tr>')
        ability_html = (
            f'<p>实际年龄：<b>{esc_h(chron)}</b>（按评估日与出生日期计算）；'
            f'综合能力年龄（取各维度最低值为下限）：<b>{esc_h(overall)}</b>。</p>'
            f'<p class="muted">判定方法：得分率 ≥ {thr}% 为基本掌握，&lt;{thr}% 为临界；'
            f'取连续掌握最远点线性插值定龄；±3 个月内为符合年龄，超出为超前/滞后。综合能力年龄取各维度最低值。</p>'
            '<table class="grid"><thead><tr><th>领域</th><th>能力年龄</th><th>实际年龄</th>'
            '<th>结论</th><th>基本掌握/定龄项</th><th>判定说明</th></tr></thead><tbody>'
            + ab_rows + '</tbody></table>'
        )
    else:
        ability_html = "<p>暂无能力年龄数据。</p>"

    return TEMPLATE_HTML.format(
        name=esc_h(client.get("name", "")),
        date=esc_h(assessment.get("date", "")),
        info_rows=info_rows,
        total_score=total["score"], total_max=total["max_total"], total_rate=total["rate"],
        total_rated=total["rated"], total_items=total["total_items"],
        detail_rows=detail_rows,
        chart_imgs=chart_imgs,
        analysis_html=analysis_html,
        sugg_html=sugg_html,
        ability_html=ability_html,
    )


# 报告 HTML 预览模板（含打印样式）
TEMPLATE_HTML = """<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>评估报告 · {name}</title>
<style>
  * {{ box-sizing: border-box; }}
  body {{ font-family: -apple-system,"PingFang SC","Microsoft YaHei",sans-serif;
         background:#eef1f5; color:#1f2933; margin:0; padding:24px; }}
  .page {{ max-width: 920px; margin:0 auto; background:#fff; padding:32px 40px;
           border-radius:10px; box-shadow:0 2px 12px rgba(0,0,0,.08); }}
  .toolbar {{ max-width:920px; margin:0 auto 14px; display:flex; gap:10px; }}
  .toolbar button {{ background:#2f6fed; color:#fff; border:none; border-radius:7px;
                     padding:8px 14px; cursor:pointer; font-size:13px; }}
  .toolbar button:hover {{ background:#1f54c4; }}
  h1 {{ text-align:center; font-size:22px; color:#2f6fed; margin:0 0 4px; }}
  h2 {{ font-size:17px; color:#1f54c4; border-left:4px solid #2f6fed; padding-left:10px;
        margin:26px 0 12px; }}
  h3 {{ font-size:15px; color:#334155; margin:14px 0 8px; }}
  h4 {{ font-size:14px; color:#1f2933; margin:16px 0 8px; background:#eef3ff;
        padding:6px 10px; border-radius:6px; }}
  p {{ line-height:1.8; margin:8px 0; font-size:13.5px; }}
  .muted {{ color:#66727f; font-size:12.5px; }}
  table.grid {{ width:100%; border-collapse:collapse; font-size:12.5px; margin:10px 0; }}
  table.grid th {{ background:#2f6fed; color:#fff; padding:7px 8px; text-align:left; }}
  table.grid td {{ border:1px solid #cbd5e1; padding:6px 8px; vertical-align:top; }}
  table.grid tbody tr:nth-child(even) {{ background:#f5f7fa; }}
  table.info {{ width:100%; border-collapse:collapse; font-size:13px; }}
  table.info th {{ background:#eef3ff; text-align:left; padding:7px 10px; width:120px;
                  border:1px solid #cbd5e1; color:#1f54c4; }}
  table.info td {{ border:1px solid #cbd5e1; padding:7px 10px; }}
  .overview {{ display:flex; gap:18px; flex-wrap:wrap; align-items:center; margin:10px 0; }}
  .overview .big {{ font-size:30px; font-weight:700; color:#2f6fed; }}
  .overview .big small {{ font-size:14px; color:#66727f; font-weight:400; }}
  .pill {{ display:inline-block; background:#eef3ff; color:#1f54c4; padding:3px 10px;
           border-radius:12px; font-size:12.5px; }}
  .chart {{ text-align:center; margin:10px 0; }}
  .chart img {{ max-width:100%; border:1px solid #e3e8ef; border-radius:8px; }}
  .chart .cap {{ color:#66727f; font-size:12px; margin-top:6px; }}
  .sugg-block {{ margin:8px 0 4px; }}
  @media print {{ body {{ background:#fff; padding:0; }} .toolbar {{ display:none; }}
    .page {{ box-shadow:none; max-width:none; }} }}
</style></head>
<body>
  <div class="toolbar">
    <button onclick="window.print()">🖨 打印 / 另存为 PDF</button>
    <button onclick="history.back()">← 返回评估系统</button>
  </div>
  <div class="page">
    <h1>语言与学习技能评估报告</h1>
    <table class="info">{info_rows}</table>
    <h2>一、总体得分概览</h2>
    <div class="overview">
      <div class="big">{total_score}<small> / {total_max} 分</small></div>
      <div><span class="pill">得分率 {total_rate}%</span>
           <span class="pill">已评 {total_rated}/{total_items} 项</span></div>
    </div>
    <h2>二、各维度得分明细</h2>
    <table class="grid"><thead><tr><th>领域</th><th>得分</th><th>满分</th>
      <th>得分率</th><th>已评/总项</th></tr></thead><tbody>{detail_rows}</tbody></table>
    <h2>三、各维度得分率可视化</h2>
    {chart_imgs}
    <h2>四、图表解读</h2>
    {analysis_html}
    <h2>五、干预建议（基于本次评估）</h2>
    {sugg_html}
    <h2>六、各维度能力年龄评估</h2>
    {ability_html}
  </div>
</body></html>"""
